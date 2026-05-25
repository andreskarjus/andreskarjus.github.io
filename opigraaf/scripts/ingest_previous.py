from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Iterable

import fitz
from docx import Document

from common import ROOT, clean_label, ensure_dirs, safe_filename_from_url, sha1_text, write_csv, write_jsonl

WORKSPACE = ROOT.parent
PREVIOUS_DIR = WORKSPACE / "previous"


def repair_mojibake(text: str) -> str:
    # Prior-agent text files appear to contain UTF-8 decoded as Latin-1 in places.
    # This reversible repair fixes common Estonian mojibake without touching good text.
    if not any(mark in text for mark in ["Ã", "â", "Å"]):
        return text
    try:
        fixed = text.encode("latin1", errors="ignore").decode("utf-8", errors="ignore")
        if fixed.count("õ") + fixed.count("ü") + fixed.count("ä") + fixed.count("ö") > text.count("õ") + text.count("ü") + text.count("ä") + text.count("ö"):
            return fixed
    except Exception:
        pass
    replacements = {
        "Ãµ": "õ",
        "Ã•": "Õ",
        "Ã¤": "ä",
        "Ã„": "Ä",
        "Ã¶": "ö",
        "Ã–": "Ö",
        "Ã¼": "ü",
        "Ãœ": "Ü",
        "Å¡": "š",
        "Å ": "Š",
        "Å¾": "ž",
        "Å½": "Ž",
        "â€“": "-",
        "â†’": "->",
        "â€œ": "\"",
        "â€": "\"",
        "â€ž": "\"",
        "â€™": "'",
        "Â": "",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text


def classify_file(path: Path) -> str:
    name = path.name.lower()
    if "guidance" in name or "õppekavade õpitulemitest" in name or "pilot agent results" in name:
        return "guidance"
    if "lisa-1" in name or "ainekava" in name:
        return "curriculum"
    if "hindamismudel" in name:
        return "assessment_model"
    if "töökava" in name or "iii-kooliaste" in name or "põhimõisted" in name or "häälestus" in name:
        return "local_material"
    return "local_material"


def extract_txt(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ["utf-8", "utf-8-sig", "cp1257", "cp1252"]:
        try:
            return repair_mojibake(raw.decode(enc))
        except UnicodeDecodeError:
            continue
    return repair_mojibake(raw.decode("utf-8", errors="ignore"))


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return repair_mojibake("\n".join(parts))


def extract_pdf(path: Path) -> str:
    doc = fitz.open(str(path))
    parts = []
    for page_no, page in enumerate(doc, start=1):
        text = page.get_text("text").strip()
        if text:
            parts.append(f"[page {page_no}]\n{text}")
    return repair_mojibake("\n\n".join(parts))


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        return extract_txt(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    return ""


def chunk_text(text: str, max_chars: int = 1800, overlap: int = 180) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks = []
    buf = ""
    for para in paras:
        if len(buf) + len(para) + 2 <= max_chars:
            buf = (buf + "\n\n" + para).strip()
        else:
            if buf:
                chunks.append(buf)
            if len(para) <= max_chars:
                buf = para
            else:
                for i in range(0, len(para), max_chars - overlap):
                    chunks.append(para[i : i + max_chars])
                buf = ""
    if buf:
        chunks.append(buf)
    return chunks


def infer_subjects(text: str, path: Path) -> list[str]:
    hay = (path.name + " " + text[:2000]).lower()
    subjects = []
    if "eesti keel" in hay or "keele" in hay:
        subjects.append("Eesti keel")
    if "kirjandus" in hay:
        subjects.append("Kirjandus")
    return subjects or ["Eesti keel"]


def infer_stage(text: str, path: Path) -> str | None:
    name = path.name.lower()
    hay = (path.name + " " + text[:3000]).lower()
    # Filename-level scope is usually more reliable than broad references inside
    # curriculum/guidance text, which may mention multiple stages.
    if "iii-kooliaste" in name or "iii kooliaste" in name or "9." in name or "ix" in name:
        return "III kooliaste"
    if "gümnaas" in name or "keskharidus" in name or "keskkool" in name:
        return "Gümnaasium"
    if "iii kooliaste" in hay or "9. klass" in hay or "ix klass" in hay:
        return "III kooliaste"
    if "gümnaas" in hay or "keskharidus" in hay or "keskkool" in hay:
        return "Gümnaasium"
    if "põhikool" in hay or "põhiharidus" in hay:
        return "Põhikool"
    return None


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--previous-dir", default=str(PREVIOUS_DIR))
    args = parser.parse_args()
    ensure_dirs()
    previous_dir = Path(args.previous_dir)
    raw_dir = ROOT / "data/raw/previous"
    raw_dir.mkdir(parents=True, exist_ok=True)

    inventory = []
    chunks = []
    guidance_notes = []
    for path in sorted(previous_dir.iterdir()):
        if not path.is_file():
            continue
        kind = classify_file(path)
        target = raw_dir / path.name
        if not target.exists() or target.stat().st_size != path.stat().st_size:
            shutil.copy2(path, target)
        try:
            text = extract_text(path)
            status = "ok" if text.strip() else "empty_or_unsupported"
            error = ""
        except Exception as exc:  # noqa: BLE001
            text = ""
            status = "error"
            error = repr(exc)
        subjects = infer_subjects(text, path)
        stage = infer_stage(text, path)
        source_url = str(path.resolve())
        file_chunks = chunk_text(text)
        for idx, chunk in enumerate(file_chunks, start=1):
            chunk_id = f"previous:{sha1_text(source_url + str(idx) + chunk, 20)}"
            row = {
                "id": chunk_id,
                "source_system": "previous_local",
                "source_url": source_url,
                "source_file": path.name,
                "source_kind": kind,
                "subjects": subjects,
                "subject": subjects[0] if subjects else None,
                "school_stage": stage,
                "heading_path": [kind, path.name],
                "text": chunk,
                "method": "local_file_text_extraction",
            }
            chunks.append(row)
        if kind == "guidance":
            guidance_notes.append({"file": path.name, "source_url": source_url, "text_preview": text[:4000]})
        inventory.append(
            {
                "file": path.name,
                "kind": kind,
                "extension": path.suffix.lower(),
                "bytes": path.stat().st_size,
                "status": status,
                "characters": len(text),
                "chunks": len(file_chunks),
                "subjects": "|".join(subjects),
                "school_stage": stage,
                "raw_copy": str(target),
                "error": error,
            }
        )

    write_csv(ROOT / "reports/previous_source_inventory.csv", inventory)
    write_jsonl(ROOT / "data/interim/previous/chunks.jsonl", chunks)
    write_jsonl(ROOT / "data/interim/previous/guidance_notes.jsonl", guidance_notes)

    curriculum_chunks = [c for c in chunks if c["source_kind"] in {"curriculum", "assessment_model", "guidance"}]
    material_chunks = [c for c in chunks if c["source_kind"] == "local_material"]
    write_jsonl(ROOT / "data/interim/curriculum/previous_chunks.jsonl", curriculum_chunks)
    write_jsonl(ROOT / "data/interim/materials/previous_material_chunks.jsonl", material_chunks)

    md = [
        "# Previous Source Bundle Ingestion",
        "",
        f"- Files scanned: {len(inventory)}",
        f"- Text chunks extracted: {len(chunks)}",
        f"- Curriculum/guidance/assessment chunks: {len(curriculum_chunks)}",
        f"- Local material chunks: {len(material_chunks)}",
        "",
        "## Useful Guidance Absorbed",
        "",
        "- Use reusable type/subtype units with parameters rather than creating separate duplicate nodes per grade or genre.",
        "- Keep task/criterion taxonomies teacher-usable and grounded in official curriculum/exam practice.",
        "- For milestone 1, this guidance mainly informs granularity and future task/criterion extensions.",
        "",
        "## Keskkool/Gümnaasium Gap",
        "",
        "- The local bundle is mostly põhikool/III kooliaste material.",
        "- Comparable gümnaasium/keskharidus material should be harvested from official oppekava/Riigi Teataja/oppekava.ee sources before drawing coverage conclusions.",
    ]
    (ROOT / "reports/previous_ingestion_report.md").write_text("\n".join(md) + "\n", encoding="utf-8")
    print(f"Ingested {len(inventory)} previous files into {len(chunks)} chunks")


if __name__ == "__main__":
    main()
